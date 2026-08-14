#!/usr/bin/env python3
"""
===============================================================================
FastAPI Cloud HTTP Service Unit Tests
===============================================================================
Tests:
  1. GET  /health returns 200 with service metadata.
  2. POST /redact with a valid synthetic DOCX returns 200 and a valid redacted DOCX.
  3. POST /redact with a non-DOCX upload is rejected with HTTP 400.
  4. POST /redact with empty or corrupted payloads is safely rejected.
  5. API responses and error messages contain zero raw PII exposure.
===============================================================================
"""

import io
import unittest
import docx
from fastapi.testclient import TestClient

from app import app


class TestPIIRedactionAPI(unittest.TestCase):
    @classmethod
    def setUpClass(cls):
        cls.client = TestClient(app)

    def test_00_root_landing_page(self):
        """GET / returns 200 HTML with self-explanatory landing page documentation."""
        response = self.client.get("/")
        self.assertEqual(response.status_code, 200)
        self.assertIn("text/html", response.headers.get("content-type", ""))
        self.assertIn("PII Redaction Service", response.text)
        self.assertIn("/docs", response.text)
        self.assertIn("/redact", response.text)
        self.assertIn("Microsoft Word", response.text)

    def test_01_health_check(self):
        """GET /health returns 200 and expected metadata schema."""
        response = self.client.get("/health")
        self.assertEqual(response.status_code, 200)
        data = response.json()
        self.assertEqual(data.get("status"), "healthy")
        self.assertEqual(data.get("service"), "pii-redaction-service")
        self.assertEqual(data.get("version"), "2.2.0")

    def test_02_redact_valid_docx(self):
        """POST /redact with a valid DOCX processes PII and returns a valid DOCX file."""
        # Create a synthetic in-memory DOCX document with test PII
        doc = docx.Document()
        doc.add_paragraph("Director: Kushal Hegde and Email: cs.connect@kshinternational.com")
        doc.add_paragraph("Phone: +91 98765 43210 and SSN: 123-45-6789")

        doc_buf = io.BytesIO()
        doc.save(doc_buf)
        doc_buf.seek(0)

        response = self.client.post(
            "/redact",
            files={"file": ("test_sample.docx", doc_buf.getvalue(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
        )

        self.assertEqual(response.status_code, 200)
        self.assertIn(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            response.headers.get("content-type", "")
        )

        # Verify returned byte stream is a valid DOCX
        out_doc = docx.Document(io.BytesIO(response.content))
        full_text = "\n".join(p.text for p in out_doc.paragraphs)

        # Verify sensitive PII strings are redacted
        self.assertNotIn("Kushal Hegde", full_text)
        self.assertNotIn("cs.connect@kshinternational.com", full_text)
        self.assertNotIn("+91 98765 43210", full_text)
        self.assertNotIn("123-45-6789", full_text)

    def test_03_reject_non_docx_file(self):
        """POST /redact rejects files with invalid extensions."""
        response = self.client.post(
            "/redact",
            files={"file": ("document.txt", b"plain text data", "text/plain")}
        )
        self.assertEqual(response.status_code, 400)
        self.assertIn("detail", response.json())

    def test_04_reject_corrupted_docx(self):
        """POST /redact rejects fake or corrupted .docx files."""
        response = self.client.post(
            "/redact",
            files={"file": ("corrupt.docx", b"NOT_A_VALID_ZIP_PACKAGE", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
        )
        self.assertEqual(response.status_code, 400)
        self.assertIn("detail", response.json())

    def test_05_reject_empty_file(self):
        """POST /redact rejects empty payload."""
        response = self.client.post(
            "/redact",
            files={"file": ("empty.docx", b"", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
        )
        self.assertEqual(response.status_code, 400)

    def test_06_zero_raw_pii_in_api_errors(self):
        """Verify API error messages do not reflect raw sensitive inputs."""
        sensitive_payload = b"John Doe secret payload"
        response = self.client.post(
            "/redact",
            files={"file": ("test.pdf", sensitive_payload, "application/pdf")}
        )
        self.assertEqual(response.status_code, 400)
        self.assertNotIn("John Doe", response.text)
        self.assertNotIn("secret payload", response.text)


if __name__ == "__main__":
    unittest.main()
