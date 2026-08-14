#!/usr/bin/env python3
"""
===============================================================================
Automated Post-Redaction Leakage Validation Unit Tests
===============================================================================
Description:
    Validates independent PII Leakage Scanner:
    1. Correctly redacted document passes with status 'PASS' and 0 residual PII.
    2. Unredacted document containing raw PII fails with status 'FAIL'.
===============================================================================
"""

import os
import sys
import unittest
import docx

sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), "..")))
from pii_redactor import PIIDetector, PIILeakageValidator, DomainProfile

class TestLeakageValidation(unittest.TestCase):
    @classmethod
    def setUpClass(cls):
        cls.base_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
        cls.redacted_docx_path = os.path.join(cls.base_dir, "Red Herring Prospectus_redacted.docx")
        cls.gt_path = os.path.join(cls.base_dir, "ground_truth.json")
        
        cls.detector = PIIDetector(method="hybrid", domain_profile=DomainProfile.get_rhp_default_profile())
        cls.validator = PIILeakageValidator(detector=cls.detector, ground_truth_path=cls.gt_path)

    def test_01_redacted_document_passes_validation(self):
        """Verify that the properly redacted document passes validation with 0 residual PII."""
        self.assertTrue(os.path.exists(self.redacted_docx_path), "Redacted docx file missing.")
        
        report = self.validator.validate_document(self.redacted_docx_path)
        
        self.assertEqual(report["status"], "PASS")
        self.assertEqual(report["total_residual_entities"], 0)
        self.assertIn("Paragraphs, Tables", report["validation_scope"])

    def test_02_unredacted_document_fails_validation(self):
        """Verify that a document containing raw unredacted PII fails validation."""
        temp_docx_path = os.path.join(self.base_dir, "scratch", "temp_unredacted_test.docx")
        os.makedirs(os.path.dirname(temp_docx_path), exist_ok=True)
        
        doc = docx.Document()
        p = doc.add_paragraph()
        p.add_run("Unredacted email is cs.connect@kshinternational.com and SSN 123-45-6789.")
        doc.save(temp_docx_path)

        report = self.validator.validate_document(temp_docx_path)
        
        self.assertEqual(report["status"], "FAIL")
        self.assertGreater(report["total_residual_entities"], 0)
        
        # Cleanup temp file
        if os.path.exists(temp_docx_path):
            os.remove(temp_docx_path)

if __name__ == "__main__":
    unittest.main()
