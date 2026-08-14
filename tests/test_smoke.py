#!/usr/bin/env python3
"""
===============================================================================
End-to-End Smoke Test for DOCX Redaction Pipeline.
===============================================================================
Description:
    Simple end-to-end integration smoke test verifying full document redaction,
    pseudonymization, and post-redaction leakage validation.
===============================================================================
"""

import os
import sys
import unittest
import docx

sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), "..")))
from pii_redactor import PIIDetector, PIIAnonymizer, DocxRedactor, PIILeakageValidator

class TestEndToEndSmoke(unittest.TestCase):
    def setUp(self):
        self.base_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
        self.input_path = os.path.join(self.base_dir, "scratch", "smoke_input.docx")
        self.output_path = os.path.join(self.base_dir, "scratch", "smoke_output_redacted.docx")
        os.makedirs(os.path.dirname(self.input_path), exist_ok=True)

        doc = docx.Document()
        doc.add_heading("Smoke Test Document", level=1)
        doc.add_paragraph("Contact person: Kushal Hegde at email contact@kshinternational.com or call +91 98765 43210.")
        table = doc.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "Address:"
        table.cell(0, 1).text = "201, Tower 2, Montreal Business Centre, Off Pallod Farms, Baner Pune – 411 045 Maharashtra, India"
        doc.save(self.input_path)

    def tearDown(self):
        for p in [self.input_path, self.output_path]:
            if os.path.exists(p):
                try:
                    os.remove(p)
                except Exception:
                    pass

    def test_end_to_end_redaction_smoke(self):
        detector = PIIDetector(method="hybrid")
        anonymizer = PIIAnonymizer(strategy="synthetic", seed=42)
        redactor = DocxRedactor(detector=detector, anonymizer=anonymizer)

        summary = redactor.redact_document(self.input_path, self.output_path)

        self.assertTrue(os.path.exists(self.output_path), "Redacted docx file was not created!")
        self.assertGreater(summary["summary"]["total_pii_redacted"], 0, "No PII was redacted!")

        # Run independent leakage validation
        validator = PIILeakageValidator(detector=detector)
        val_report = validator.validate_document(self.output_path)

        self.assertEqual(val_report["status"], "PASS")
        self.assertEqual(val_report["total_residual_entities"], 0)

        # Verify raw PII strings do NOT exist in redacted document
        doc_redacted = docx.Document(self.output_path)
        redacted_text = "\n".join([p.text for p in doc_redacted.paragraphs] + [c.text for t in doc_redacted.tables for r in t.rows for c in r.cells])

        self.assertNotIn("Kushal Hegde", redacted_text)
        self.assertNotIn("contact@kshinternational.com", redacted_text)
        self.assertNotIn("98765 43210", redacted_text)

if __name__ == "__main__":
    unittest.main()
