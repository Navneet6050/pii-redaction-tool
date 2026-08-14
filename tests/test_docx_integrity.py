#!/usr/bin/env python3
"""
===============================================================================
Automated Document Integrity & Run-Aware Redaction Tests
===============================================================================
Description:
    Rigorously validates DOCX document structure integrity and run-aware styling:
    1. Single-run PII entity replacement (bold/italic styling preservation)
    2. Multi-run PII entity replacement (cross-run span mapping)
    3. Multiple PII entities in a single paragraph
    4. Table & Nested Table PII redaction
    5. Header & Footer PII redaction
    6. Structural comparison (Paragraph, Table, Section counts)
===============================================================================
"""

import os
import sys
import unittest
import docx

# Add parent directory to sys.path
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), "..")))

from pii_redactor import (
    PIIDetector, PIIAnonymizer, DocxRedactor, DomainProfile
)

class TestDocxIntegrity(unittest.TestCase):
    @classmethod
    def setUpClass(cls):
        cls.base_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
        cls.input_path = os.path.join(cls.base_dir, "Red Herring Prospectus.docx")
        cls.output_path = os.path.join(cls.base_dir, "Red Herring Prospectus_redacted.docx")
        
        cls.detector = PIIDetector(method="hybrid", domain_profile=DomainProfile.get_rhp_default_profile())
        cls.anonymizer = PIIAnonymizer(strategy="synthetic")
        cls.redactor = DocxRedactor(detector=cls.detector, anonymizer=cls.anonymizer)

    def test_01_single_run_formatting_preservation(self):
        """Test that single-run PII replacement preserves bold and italic styling."""
        doc = docx.Document()
        p = doc.add_paragraph()
        run1 = p.add_run("Contact email is ")
        run2 = p.add_run("cs.connect@kshinternational.com")
        run2.bold = True
        run2.italic = True
        run3 = p.add_run(" for inquiries.")

        self.redactor._process_paragraph_run_aware(p, "test_p1")

        self.assertNotIn("cs.connect@kshinternational.com", p.text)
        # Verify run2 formatting attributes remain intact
        self.assertTrue(p.runs[1].bold)
        self.assertTrue(p.runs[1].italic)

    def test_02_multi_run_span_redaction(self):
        """Test PII entity spanning across multiple runs."""
        doc = docx.Document()
        p = doc.add_paragraph()
        r1 = p.add_run("Email: cs.connect")
        r1.bold = True
        r2 = p.add_run("@kshinternational.com")
        r2.italic = True
        r3 = p.add_run(" for help.")

        self.redactor._process_paragraph_run_aware(p, "test_p2")

        self.assertNotIn("cs.connect@kshinternational.com", p.text)
        # Verify first run retained bold styling and original non-PII text "Email: "
        self.assertTrue(p.runs[0].bold)
        self.assertIn("Email: ", p.runs[0].text)

    def test_03_multiple_entities_in_single_paragraph(self):
        """Test replacing multiple PII entities in a single paragraph without offset drift."""
        doc = docx.Document()
        p = doc.add_paragraph("Kushal Hegde email is cs.connect@kshinternational.com or call +91 9876543210.")

        self.redactor._process_paragraph_run_aware(p, "test_p3")

        self.assertNotIn("Kushal Hegde", p.text)
        self.assertNotIn("cs.connect@kshinternational.com", p.text)
        self.assertNotIn("+91 9876543210", p.text)

    def test_04_table_and_nested_table_redaction(self):
        """Test redaction inside table cells and nested tables."""
        doc = docx.Document()
        table = doc.add_table(rows=2, cols=2)
        cell = table.cell(0, 0)
        cell.paragraphs[0].text = "Promoter: Kushal Hegde"
        
        # Add nested table inside cell
        nested_table = cell.add_table(rows=1, cols=1)
        nested_table.cell(0, 0).paragraphs[0].text = "Contact: cs.connect@kshinternational.com"

        self.redactor._process_table(table, "table_0")

        self.assertNotIn("Kushal Hegde", cell.paragraphs[0].text)
        self.assertNotIn("cs.connect@kshinternational.com", nested_table.cell(0, 0).paragraphs[0].text)

    def test_05_header_and_footer_redaction(self):
        """Test redaction in section headers and footers."""
        doc = docx.Document()
        section = doc.sections[0]
        header = section.header
        header.paragraphs[0].text = "KSH International Limited - RHP Filing"

        self.redactor._process_paragraph_run_aware(header.paragraphs[0], "header_p0")

        self.assertNotIn("KSH International Limited", header.paragraphs[0].text)

    def test_06_document_structure_integrity(self):
        """Verify before/after structural equality (Paragraphs, Tables, Sections)."""
        self.assertTrue(os.path.exists(self.input_path), "Input docx missing.")
        self.assertTrue(os.path.exists(self.output_path), "Output redacted docx missing.")

        doc_orig = docx.Document(self.input_path)
        doc_redacted = docx.Document(self.output_path)

        # Structure equality checks
        self.assertEqual(len(doc_orig.paragraphs), len(doc_redacted.paragraphs), "Paragraph counts mismatch!")
        self.assertEqual(len(doc_orig.tables), len(doc_redacted.tables), "Table counts mismatch!")
        self.assertEqual(len(doc_orig.sections), len(doc_redacted.sections), "Section counts mismatch!")
        
        # Verify row/column counts across all tables
        for t_orig, t_redacted in zip(doc_orig.tables, doc_redacted.tables):
            self.assertEqual(len(t_orig.rows), len(t_redacted.rows))

if __name__ == "__main__":
    unittest.main()
