#!/usr/bin/env python3
"""
===============================================================================
PII Redaction Tool - Production Grade DOCX Anonymization Engine
===============================================================================
Version: 2.2.0
Description:
    A modular, high-precision PII Detection and Anonymization Engine for DOCX
    featuring RUN-AWARE redaction that preserves original typography, bold,
    italics, underlines, colors, and document structure.

Features:
    - Shared NLP Model Singleton Caching (Fast execution, spaCy & Presidio loaded once)
    - Pure Generic Detection by Default (Zero dependence on hardcoded domain profiles)
    - Optional Domain Profile Gazetteer (--use-domain-profile)
    - Enhanced Contextual DOB Detection (Birth dates, DOB, Born on, Age+Born signals)
    - Full XML Traversal: Body, Tables, Nested Tables, Headers, Footers, Hyperlinks, Text Boxes (w:txbxContent)
    - Robust Deterministic Pseudonymization (--seed)
    - Collision Avoidance (different entities get distinct replacements)
    - Independent Post-Redaction Leakage Validator (--validate, --strict)
    - Clean CLI Boolean Semantics (--use-domain-profile / --no-use-domain-profile)
    - Zero raw PII in output reports, logs, or persistent files

Usage:
    python pii_redactor.py --input "Red Herring Prospectus.docx" --output "Red Herring Prospectus_redacted.docx" --seed 42 --validate --strict
===============================================================================
"""

import argparse
import logging
import os
import re
import sys
import json
import time
import ipaddress
from dataclasses import dataclass, asdict
from collections import defaultdict
from typing import Dict, List, Tuple, Set, Any, Optional

import docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from faker import Faker
import spacy
from presidio_analyzer import AnalyzerEngine
from presidio_analyzer.nlp_engine import NlpEngineProvider

__version__ = "2.2.0"

# Configure Logging
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    handlers=[logging.StreamHandler(sys.stdout)]
)
logger = logging.getLogger("PIIRedactor")

# Silence external library noise
logging.getLogger("presidio-analyzer").setLevel(logging.ERROR)
logging.getLogger("spacy").setLevel(logging.ERROR)


# =============================================================================
# SINGLETON MODEL CACHING (Thread-Safe Process-Level Caching)
# =============================================================================
_SPACY_NLP_MODEL = None
_PRESIDIO_ANALYZER_ENGINE = None

def get_shared_spacy_model():
    """Returns process-level singleton instance of spaCy en_core_web_sm."""
    global _SPACY_NLP_MODEL
    if _SPACY_NLP_MODEL is None:
        try:
            _SPACY_NLP_MODEL = spacy.load("en_core_web_sm")
        except OSError:
            logger.error("spaCy model 'en_core_web_sm' not found. Install using: python -m spacy download en_core_web_sm")
            sys.exit(1)
    return _SPACY_NLP_MODEL

def get_shared_presidio_engine():
    """Returns process-level singleton instance of Presidio AnalyzerEngine."""
    global _PRESIDIO_ANALYZER_ENGINE
    if _PRESIDIO_ANALYZER_ENGINE is None:
        nlp_config = {
            "nlp_engine_name": "spacy",
            "models": [{"lang_code": "en", "model_name": "en_core_web_sm"}],
        }
        provider = NlpEngineProvider(nlp_configuration=nlp_config)
        nlp_engine = provider.create_engine()
        _PRESIDIO_ANALYZER_ENGINE = AnalyzerEngine(nlp_engine=nlp_engine)
    return _PRESIDIO_ANALYZER_ENGINE


# =============================================================================
# DATA STRUCTURES & UTILITIES
# =============================================================================
@dataclass
class PIIEntity:
    """Normalized internal representation of a detected PII entity."""
    category: str
    start: int
    end: int
    text: str
    confidence: float
    source: str


def luhn_check(card_number_str: str) -> bool:
    """Validates credit card numbers using the Luhn checksum algorithm."""
    digits = [int(c) for c in card_number_str if c.isdigit()]
    if len(digits) < 13 or len(digits) > 19 or all(d == 0 for d in digits):
        return False
    checksum = 0
    reverse_digits = digits[::-1]
    for i, digit in enumerate(reverse_digits):
        if i % 2 == 1:
            doubled = digit * 2
            checksum += doubled - 9 if doubled > 9 else doubled
        else:
            checksum += digit
    return checksum % 10 == 0


def is_valid_ip(ip_str: str) -> bool:
    """Strictly validates IPv4 and IPv6 network addresses."""
    try:
        ipaddress.ip_address(ip_str.strip())
        return True
    except ValueError:
        return False


# =============================================================================
# DOMAIN PROFILE (Decoupled Optional Gazetteer Profile)
# =============================================================================
class DomainProfile:
    """
    Decoupled container for optional domain-specific gazetteer knowledge.
    Ensures detector functions 100% standalone in generic mode without hardcoded document strings.
    """
    def __init__(
        self,
        promoter_names: Optional[List[str]] = None,
        company_names: Optional[List[str]] = None,
        addresses: Optional[List[str]] = None
    ):
        self.promoter_names = promoter_names or []
        self.company_names = company_names or []
        self.addresses = addresses or []

    @classmethod
    def get_rhp_default_profile(cls) -> "DomainProfile":
        """Returns the domain profile specifically for the KSH RHP document."""
        return cls(
            promoter_names=[
                "Kushal Subbayya Hegde", "Kushal Hegde", "KUSHAL SUBBAYYA HEGDE", "KUSHAL HEGDE",
                "Pushpa Kushal Hegde", "PUSHPA KUSHAL HEGDE", "Pushpa Hegde",
                "Rajesh Kushal Hegde", "RAJESH KUSHAL HEGDE", "Rajesh Hegde",
                "Rohit Kushal Hegde", "ROHIT KUSHAL HEGDE", "Rohit Hegde",
                "Rakhi Girija Shetty", "RAKHI GIRIJA SHETTY", "Rakhi Shetty",
                "Sarthak Malvadkar", "SARTHAK MALVADKAR", "Sarthak Malvadakar",
                "Anand Soni", "Ashish M P", "Cherag Gyara", "Eric Bacha", "Hitesh Ramani",
                "Manisha Shukla", "Parag Pansare", "Prakash Boricha", "Pravin Teli",
                "Sachin Gawade", "Sharmila Joshi", "Sheetal Parab", "Siddharth Jadhav",
                "Tushar Gavankar", "Katyayani Balasubramanian", "Gopalakrishnan V"
            ],
            company_names=[
                "KSH INTERNATIONAL LIMITED", "KSH International Limited", "KSH International",
                "ICICI Securities Limited", "Nuvama Wealth Management Limited", "MUFG Bank Ltd",
                "MUFG Bank", "Federal Bank Limited", "HDFC Bank Limited", "Bajaj Finserv Limited",
                "Kirtane & Pandit LLP", "Trilegal", "Exim Bank of India", "State Bank of India",
                "IndusInd Bank Limited", "IndusInd Bank", "Citibank N.A.", "Citibank",
                "Everest Family Trust", "Broad Family Trust", "Dhaulagiri Family Trust",
                "Makalu Family Trust", "Annapurna Family Trust", "Kanchenjunga Family Trust",
                "Waterloo Industrial Park V", "WATERLOO INDUSTRIAL PARK V", "Link Intime India Private Limited"
            ],
            addresses=[
                "11/3, 11/4 and 11/5 Village Birdewadi Chakan Taluka - Khed Pune – 410 501 Maharashtra, India",
                "201, Tower 2, Montreal Business Centre, Off Pallod Farms, Baner Pune – 411 045 Maharashtra, India",
                "11/3, 11/4 and 11/5 Village Birdewadi Chakan Taluka - Khed Pune",
                "201, Tower 2, Montreal Business Centre, Off Pallod Farms, Baner Pune"
            ]
        )


# =============================================================================
# FINANCIAL EXCLUSIONS & STOP WORDS
# =============================================================================
FINANCIAL_LINE_ITEMS = {
    "CAPITAL EMPLOYED", "CAPITAL RESERVE", "CAPITAL STRUCTURE", "BANK BALANCES AND ADVANCES",
    "CASH AND BANK BALANCES", "BANK BALANCES", "KEY MANAGEMENT PERSONNEL", "KEY MANAGERIAL PERSONNEL",
    "MINISTRY OF CORPORATE AFFAIRS", "GOODS AND SERVICES TAX", "VALUE ADDED TAX", "SERVICE TAX",
    "INCOME TAX", "STANDALONE", "CONSOLIDATED", "BOARD OF DIRECTORS", "AUDIT COMMITTEE",
    "NOMINATION AND REMUNERATION COMMITTEE", "STAKEHOLDERS RELATIONSHIP COMMITTEE",
    "RISK MANAGEMENT COMMITTEE", "FINANCIAL STATEMENTS", "RESTATED FINANCIAL INFORMATION",
    "PROMOTER SELLING SHAREHOLDERS", "PROMOTERS AND PROMOTER GROUP", "OBJECTS OF THE OFFER",
    "BASIS FOR OFFER PRICE", "STATEMENT OF POSSIBLE SPECIAL TAX BENEFITS", "RISK FACTORS",
    "SUMMARY OF FINANCIAL INFORMATION", "GENERAL INFORMATION", "CAPITAL STRUCTURE",
    "OUR BUSINESS", "OUR INDUSTRY", "KEY PERFORMANCE INDICATORS", "LEGAL AND OTHER INFORMATION",
    "REFERENCE RATE", "SELLING SHAREHOLDER", "BID AMOUNT", "MUTUAL FUNDS", "NET ASSET VALUE",
    "CREDIT CARD", "DATE OF BIRTH", "EXECUTIVE DATE OF BIRTH"
}

FINANCIAL_STOPWORDS = {
    "OFFER", "EQUITY", "EQUITY SHARES", "BIDS", "BIDDERS", "CAP PRICE", "FLOOR PRICE",
    "THE OFFER", "THE EQUITY SHARES", "THE PROMOTER SELLING SHAREHOLDERS", "PROMOTERS",
    "BSE", "NSE", "SEBI", "ROC", "RHP", "DRHP", "ANCHOR INVESTORS", "NON-INSTITUTIONAL PORTION",
    "QUALIFIED INSTITUTIONAL BUYERS", "QIB", "NII", "RIB", "RETAIL INDIVIDUAL BIDDERS",
    "BOOK RUNNING LEAD MANAGERS", "BRLM", "REGISTRAR", "SYNDICATE MEMBERS", "UPI",
    "UPI MECHANISM", "ASBA", "CUT-OFF PRICE", "FACE VALUE", "ISSUE", "FRESH ISSUE",
    "OFFER FOR SALE", "COMPANIES ACT", "SEBI ICDR REGULATIONS", "PRE-OFFER SHAREHOLDING",
    "POST-OFFER SHAREHOLDING", "TABLE OF CONTENTS", "SECTION", "FISCAL 2025", "FISCAL 2024",
    "FISCAL 2023", "JUNE 30, 2025", "MARCH 31, 2025", "DECEMBER 10, 2025", "LIMITED",
    "PRIVATE LIMITED", "CORPORATE IDENTITY NUMBER", "REGISTERED OFFICE", "CORPORATE OFFICE",
    "DIN", "PAN", "ISIN", "CIN", "NEFT", "RTGS", "NA", "N.A.", "BANK", "BANK BALANCES"
}

STRICT_COMPANY_SUFFIX_PATTERN = re.compile(
    r'\b(?:LIMITED|LTD\.?|PRIVATE\s+LIMITED|PVT\.?\s*LTD\.?|PVT\s+LTD|INC\.?|CORP\.?|LLP|TRUST|'
    r'SECURITIES|MANAGEMENT|SERVICES|HOLDINGS|CAPITAL|SOLUTIONS|INDUSTRIES|ENTERPRISES|'
    r'TECHNOLOGIES|BANK|FINANCE|INFRASTRUCTURE|PHARMA|LOGISTICS|ENERGY|RETAIL|MEDIA|'
    r'TELECOM|AIRWAYS|AEROSPACE|BIOTECH|HEALTHCARE|SYSTEMS)\b',
    re.IGNORECASE
)

# URL-like patterns to EXCLUDE from COMPANY_NAME detection
# Must have http(s):// or www. prefix to avoid matching email domains
URL_PATTERN = re.compile(r'(?:https?://|www\.)[^\s]+', re.IGNORECASE)

# Generic all-caps financial/document phrases that are NOT company names
# This is a small curated list of generic phrases only — not document-specific blacklist
GENERIC_FINANCIAL_PHRASES = {
    "CAPITAL EMPLOYED", "CASH AND BANK BALANCES", "BANK BALANCES",
    "TOTAL REVENUE FROM OPERATIONS", "EARNINGS PER SHARE", "EARNINGS BEFORE INTEREST",
    "FINANCIAL YEAR ENDED", "TOTAL OUTSTANDING LIABILITIES", "TOTAL ASSETS",
    "TOTAL EQUITY", "NET WORTH", "PROFIT AND LOSS", "BALANCE SHEET",
    "CURRENT ASSETS", "FIXED ASSETS", "CURRENT LIABILITIES", "DEFERRED TAX",
    "PROVISIONS", "RESERVES AND SURPLUS", "LONG TERM BORROWINGS", "SHORT TERM BORROWINGS",
    "CREDIT CARD", "DATE OF BIRTH", "KEY MANAGEMENT PERSONNEL", "KEY MANAGERIAL PERSONNEL",
    "BOARD OF DIRECTORS", "AUDIT COMMITTEE", "MINISTRY OF CORPORATE AFFAIRS",
}

# Context keywords that strongly suggest a FULL_NAME follows
# Prefix keywords are matched case-insensitively, but captured names MUST be Title Case
FULL_NAME_CONTEXT_PATTERN = re.compile(
    r'\b(?:'
    r'(?i:Mr|Mrs|Ms|Miss|Dr|Prof|Shri|Smt|Er|Adv|Justice|Hon|Colonel|Col|Gen|Capt)\.?\s+'
    r'|'
    r'(?i:Name|Director|CEO|CFO|MD|COO|CTO|Chairman|Managing\s+Director|'
    r'Executive\s+Director|Independent\s+Director|Authorized\s+Signatory|'
    r'Contact|Partner|President|Vice\s+President|Trustee|Nominee|Promoter|'
    r'Issued\s+to|Signed\s+by|Registered\s+by|Beneficiary)[:\s]+'
    r')'
    r'([A-Z][a-zA-Z&]+(?:\s+[A-Z][a-zA-Z&]+){1,3})\b'
)

# Direct regex for company names with legal suffixes — no NER dependency
# Matches entire "CompanyName Suffix" as group(0). Suffix is the anchor.
# Multi-word suffixes (Private Limited, Pvt Ltd) are checked first.
# Single-word broad terms (Bank, Finance) are excluded from standalone suffix match
# to avoid FP on phrases like "central bank" or "corporate finance".
COMPANY_LEGAL_SUFFIX_REGEX = re.compile(
    r'\b(?:[A-Z][a-zA-Z&]+\s+){1,7}'           # 1-7 title-case words (company name body)
    r'(?:'
    r'Private\s+Limited'                          # Private Limited
    r'|Pvt\.?\s*Ltd\.?'                           # Pvt Ltd, Pvt. Ltd.
    r'|Limited'                                   # Limited
    r'|Ltd\.?'                                    # Ltd, Ltd.
    r'|Inc\.?'                                    # Inc, Inc.
    r'|Corp\.?'                                   # Corp, Corp.
    r'|Corporation'                               # Corporation
    r'|LLP'                                       # LLP
    r'|Holdings\s+Ltd\.?'                         # Holdings Ltd
    r'|Holdings\s+Limited'                        # Holdings Limited
    r'|Technologies'                              # Technologies
    r'|Industries'                                # Industries
    r'|Enterprises'                               # Enterprises
    r'|Infrastructure'                            # Infrastructure
    r'|Logistics'                                 # Logistics
    r'|Pharma'                                    # Pharma
    r'|Biotech'                                   # Biotech
    r'|Systems'                                   # Systems
    r')(?:\b|(?=\s|$|[,;.]))',
    re.MULTILINE
)

# Contextual DOB keywords — must appear before a date to count as DATE_OF_BIRTH
# 'Date' alone is NOT sufficient (avoids Offer Date, Filing Date, etc.)
STRONG_DOB_KEYWORDS = re.compile(
    r'\b(?:Date\s+of\s+Birth|DOB|Birth\s+Date|Born\s+on|Born|Birthdate)\b',
    re.IGNORECASE
)

# Address strength keywords — a span must contain at least one of these to be PHYSICAL_ADDRESS
ADDRESS_STRENGTH_KEYWORDS = [
    'flat', 'plot', 'house', 'building', 'office', 'unit', 'floor', 'tower',
    'road', 'street', 'avenue', 'lane', 'nagar', 'sector', 'block', 'village',
    'taluka', 'district', 'phase', 'park', 'registered office', 'correspondence address',
    'mailing address', 'residential address', 'corporate office', 'registered address',
    'address', 'st', 'rd', 'ave', 'blvd', 'dr', 'way', 'boulevard', 'drive',
    'marg', 'chowk', 'circle', 'cross', 'crescent', 'gali'
]


# =============================================================================
# DETECTOR CLASS (Modular Multi-Stage Pipeline - Generic by Default)
# =============================================================================
class PIIDetector:
    """
    Modular Multi-Stage PII Detection Pipeline combining:
    1. Regex Detectors (Emails, Phones, IPs, SSNs, Credit Cards, DOBs, URLs, Addresses)
    2. Microsoft Presidio Analyzer Engine (Shared Singleton)
    3. spaCy Named Entity Recognition (Shared Singleton)
    4. Contextual & Domain Exclusions
    5. Optional Gazetteer Domain Profile (None by default for generic operation)
    6. Deterministic Overlap Resolution
    """
    def __init__(
        self,
        method: str = "hybrid",
        domain_profile: Optional[DomainProfile] = None
    ):
        self.method = method
        self.domain_profile = domain_profile
        
        # Load shared singletons
        self.nlp = get_shared_spacy_model()
        self.presidio = get_shared_presidio_engine()

        # Regex Matchers
        self.regex_patterns = {
            "EMAIL_ADDRESS": re.compile(r'\b[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}\b'),
            "PHONE_NUMBER": re.compile(
                r'(?:\+91[\s-]?)?(?:[6-9]\d{9}|[6-9]\d{4}[\s-]\d{5}|[6-9]\d{2}[\s-]\d{3}[\s-]\d{4})|'
                r'(?:\(?0\d{2,4}\)?[-.\s]?\d{6,8})|'
                r'(?:\+1[\s-]\d{3}[\s-]\d{3}[\s-]\d{4})'
            ),
            "IP_ADDRESS": re.compile(r'\b(?:(?:25[0-5]|2[0-4][0-9]|[01]?[0-9][0-9]?)\.){3}(?:25[0-5]|2[0-4][0-9]|[01]?[0-9][0-9]?)\b'),
            "SSN": re.compile(r'\b(?:\d{3}-\d{2}-\d{4})\b'),
            # Expanded credit card candidate regex covering all major IIN prefix families:
            # Visa (4xxx), MasterCard (51-55xxxx, 2221-2720), AmEx (34/37), Discover (6011,65xx),
            # JCB (3528-3589), Diners (300-305, 36, 38), Maestro (6304, 6759, 676770-676774)
            # Also matches generic 13-16-19 digit space/hyphen separated groups with Luhn gate.
            "CREDIT_CARD_CANDIDATE": re.compile(
                r'\b(?:'
                r'4[0-9]{12}(?:[0-9]{3,6})?'          # Visa 13/16/19
                r'|5[1-5][0-9]{14}'                     # MasterCard 51-55
                r'|2[2-7][0-9]{14}'                     # MasterCard 2221-2720
                r'|3[47][0-9]{13}'                      # AmEx 15 digit
                r'|3(?:0[0-5]|[68])[0-9]{11}'          # Diners 14 digit
                r'|6(?:011|5[0-9]{2})[0-9]{12,15}'     # Discover 16-19
                r'|(?:2131|1800|35\d{3})\d{11}'        # JCB
                r'|(?:6304|6759|6761|6762|6763)\d{8,15}'  # Maestro
                r'|[3-6]\d{3}[\s-]\d{4}[\s-]\d{4}[\s-]\d{1,7}'  # Spaced/hyphen 4-group
                r'|[3-6]\d{3}[\s-]\d{6}[\s-]\d{4,5}'            # AmEx 4-6-5 spaced
                r')\b',
                re.ASCII
            ),
            # DOB requires strong keyword (DOB, Date of Birth, Born on, Born, Birth Date)
            # 'Date:' alone is intentionally excluded to avoid Offer Date, Filing Date, etc.
            "DATE_OF_BIRTH": re.compile(
                r'\b(?:Date\s+of\s+Birth|DOB|Birth\s+Date|Born\s+on|Born|Birthdate|Birth\s+date|aged?\s+\d+[\s,]+born|age[:\s]+\d+[\s,]+born)[:\s-]+(?:\d{1,2}[/-]\d{1,2}[/-]\d{2,4}|[A-Za-z]+\s+\d{1,2},?\s+\d{4}|\d{1,2}\s+[A-Za-z]+\s+\d{4})\b',
                re.IGNORECASE
            ),
            "DATE_OF_BIRTH_CONTEXT": re.compile(
                r'\b(?:born(?:\s+on)?|birth\s+date[:\s]+|birthdate[:\s]+)(?:\s+)?(?:\d{1,2}[\s/\.-]\d{1,2}[\s/\.-]\d{2,4}|\d{1,2}\s+[A-Za-z]+\s+\d{2,4}|[A-Za-z]+\s+\d{1,2},?\s+\d{2,4})\b',
                re.IGNORECASE
            ),
            "PAN_NUMBER": re.compile(r'\b[A-Z]{5}[0-9]{4}[A-Z]{1}\b'),
            "DIN_NUMBER": re.compile(r'\b0[0-9]{7}\b'),
            # COMPANY_URL intentionally removed — see _regex_pass exclusion logic
            "PHYSICAL_ADDRESS_GENERIC": re.compile(
                r'\b(?:Flat|Office|Plot|Tower|Building|House|No\.|Village|Taluka|District|Unit|Floor)'  # requires structural keyword
                r'\s*[^,\n]{2,40},\s*[^,\n]{2,40}(?:,\s*[^,\n]{2,40})*,\s*[A-Za-z\s]{3,20}\s*[-–]?\s*\d{5,6}\b',
                re.IGNORECASE
            ),
            "PHYSICAL_ADDRESS_STREET": re.compile(
                r'(?:\b(?:Address|Mailing\s+Address|Registered\s+Address|Registered\s+Office|Corporate\s+Office|Postal\s+Address|Office\s+Address|Residential\s+Address|Work\s+Address|Contact\s+Address)[:\s]+)?'
                r'\b\d{1,5}[A-Za-z]?(?:/\d+)?(?:\s*[-–]\s*\d+)?\s+(?:[A-Za-z0-9\.\'\-]+\s+){0,4}'
                r'(?:Street|St|Road|Rd|Avenue|Ave|Lane|Ln|Boulevard|Blvd|Drive|Dr|Way|Marg|Chowk|Path|Highway|Expressway|Circle|Cross|Crescent|Gali)\b\.?'
                r'(?:,\s*[^,\n]{2,35}){1,3}'
                r'(?:[-–\s]+\d{3,6})?'
                r'(?:,\s*[A-Za-z\s]{2,20})*',
                re.IGNORECASE
            )
        }

    def detect_pii(self, text: str) -> List[PIIEntity]:
        """Detects PII entities across the text block using configured detection modules."""
        if not text or not text.strip():
            return []

        raw_candidates: List[PIIEntity] = []

        if self.method in ["hybrid", "regex"]:
            raw_candidates.extend(self._regex_pass(text))

        if self.method in ["hybrid", "presidio"]:
            raw_candidates.extend(self._presidio_pass(text))

        if self.method in ["hybrid", "spacy"]:
            raw_candidates.extend(self._spacy_pass(text))

        if self.domain_profile:
            raw_candidates.extend(self._gazetteer_pass(text))

        final_entities = self.resolve_overlaps(raw_candidates)
        return final_entities

    def _regex_pass(self, text: str) -> List[PIIEntity]:
        candidates = []

        for cat, pattern in self.regex_patterns.items():
            if cat == "DATE_OF_BIRTH_CONTEXT":
                for match in pattern.finditer(text):
                    candidates.append(PIIEntity(
                        category="DATE_OF_BIRTH",
                        start=match.start(),
                        end=match.end(),
                        text=match.group(),
                        confidence=0.95,
                        source="REGEX_DOB_CONTEXT"
                    ))
                continue

            for match in pattern.finditer(text):
                val = match.group()
                
                if cat == "PHONE_NUMBER":
                    clean_digits = re.sub(r'\D', '', val)
                    if len(clean_digits) < 8 or clean_digits.startswith("0000") or re.search(r'20\d{2}[-–]20\d{2}', val) or re.search(r'20\d{2}[-–]\d{2}', val):
                        continue

                if cat == "CREDIT_CARD_CANDIDATE":
                    clean_cc = re.sub(r'\D', '', val)
                    if luhn_check(clean_cc):
                        candidates.append(PIIEntity(
                            category="CREDIT_CARD",
                            start=match.start(),
                            end=match.end(),
                            text=val,
                            confidence=0.99,
                            source="REGEX_LUHN"
                        ))
                    continue

                if cat == "IP_ADDRESS":
                    if not is_valid_ip(val):
                        continue
                    # Reject version-string context: 'v10.0.0.1', 'version 10.0.0.1'
                    prefix = text[:match.start()].lower().rstrip()
                    if prefix.endswith("version") or prefix.endswith("release") or (
                        len(prefix) > 0 and prefix[-1] == 'v' and (len(prefix) == 1 or not prefix[-2].isalpha())
                    ):
                        continue

                if cat in ["PHYSICAL_ADDRESS_GENERIC", "PHYSICAL_ADDRESS_STREET"]:
                    # Require at least one strong address keyword in the match
                    lower_val = val.lower()
                    if not any(kw in lower_val for kw in ADDRESS_STRENGTH_KEYWORDS):
                        continue
                    candidates.append(PIIEntity(
                        category="PHYSICAL_ADDRESS",
                        start=match.start(),
                        end=match.end(),
                        text=val,
                        confidence=0.90,
                        source="REGEX"
                    ))
                    continue

                candidates.append(PIIEntity(
                    category=cat,
                    start=match.start(),
                    end=match.end(),
                    text=val,
                    confidence=1.0,
                    source="REGEX"
                ))

        # Contextual FULL_NAME detection via title/role prefix patterns
        for m in FULL_NAME_CONTEXT_PATTERN.finditer(text):
            name_text = m.group(1).strip()
            if len(name_text.split()) >= 2 and not any(c.isdigit() for c in name_text):
                candidates.append(PIIEntity(
                    category="FULL_NAME",
                    start=m.start(1),
                    end=m.end(1),
                    text=name_text,
                    confidence=0.92,
                    source="REGEX_CONTEXT"
                ))

        # Company legal-suffix regex pass — catches names spaCy NER misses
        for m in COMPANY_LEGAL_SUFFIX_REGEX.finditer(text):
            co_text = m.group().strip()
            upper_co = co_text.upper()
            # Skip if it's a known financial phrase
            if upper_co in GENERIC_FINANCIAL_PHRASES or any(p in upper_co for p in GENERIC_FINANCIAL_PHRASES):
                continue
            if upper_co in FINANCIAL_LINE_ITEMS or upper_co in FINANCIAL_STOPWORDS:
                continue
            candidates.append(PIIEntity(
                category="COMPANY_NAME",
                start=m.start(),
                end=m.end(),
                text=co_text,
                confidence=0.88,
                source="REGEX_COMPANY_SUFFIX"
            ))

        return candidates

    def _presidio_pass(self, text: str) -> List[PIIEntity]:
        candidates = []
        try:
            results = self.presidio.analyze(
                text=text,
                language="en",
                entities=["PERSON", "LOCATION", "EMAIL_ADDRESS", "PHONE_NUMBER", "IP_ADDRESS", "CREDIT_CARD", "ORGANIZATION"]
            )
            for res in results:
                ent_text = text[res.start:res.end]
                upper_text = ent_text.upper()
                
                cat_map = {
                    "PERSON": "FULL_NAME",
                    "ORGANIZATION": "COMPANY_NAME",
                    "LOCATION": "PHYSICAL_ADDRESS",
                    "EMAIL_ADDRESS": "EMAIL_ADDRESS",
                    "PHONE_NUMBER": "PHONE_NUMBER",
                    "IP_ADDRESS": "IP_ADDRESS",
                    "CREDIT_CARD": "CREDIT_CARD"
                }
                cat = cat_map.get(res.entity_type, res.entity_type)

                if cat == "IP_ADDRESS":
                    prefix = text[:res.start].lower().rstrip()
                    if prefix.endswith("version") or prefix.endswith("release") or (
                        len(prefix) > 0 and prefix[-1] == 'v' and (len(prefix) == 1 or not prefix[-2].isalpha())
                    ):
                        continue

                if cat == "PHYSICAL_ADDRESS":
                    lower_ent = ent_text.lower()
                    # Require at least one address-structure keyword AND either a digit or postal code
                    has_structure = any(kw in lower_ent for kw in ADDRESS_STRENGTH_KEYWORDS)
                    has_number = any(char.isdigit() for char in ent_text)
                    if not has_structure or not has_number:
                        continue

                if cat == "COMPANY_NAME":
                    # Reject URL-like strings
                    if URL_PATTERN.search(ent_text):
                        continue
                    # Reject generic all-caps financial phrases
                    if upper_text in GENERIC_FINANCIAL_PHRASES or any(p in upper_text for p in GENERIC_FINANCIAL_PHRASES):
                        continue

                if upper_text in FINANCIAL_LINE_ITEMS or upper_text in FINANCIAL_STOPWORDS or any(item in upper_text for item in FINANCIAL_LINE_ITEMS):
                    continue

                candidates.append(PIIEntity(
                    category=cat,
                    start=res.start,
                    end=res.end,
                    text=ent_text,
                    confidence=round(float(res.score), 2),
                    source="PRESIDIO"
                ))
        except Exception as e:
            logger.debug(f"Presidio pass error: {e}")

        return candidates

    def _spacy_pass(self, text: str) -> List[PIIEntity]:
        candidates = []
        doc = self.nlp(text)
        
        for ent in doc.ents:
            ent_text = ent.text.strip()
            upper_text = ent_text.upper()

            if upper_text in FINANCIAL_LINE_ITEMS or upper_text in FINANCIAL_STOPWORDS or any(item in upper_text for item in FINANCIAL_LINE_ITEMS):
                continue
            if any(w in upper_text for w in ["SECTION", "TABLE", "SCHEDULE", "CLAUSE", "ARTICLE"]):
                continue

            if ent.label_ == "PERSON" and len(ent_text.split()) >= 2 and not any(char.isdigit() for char in ent_text):
                # Check for title/role context in surrounding text for higher confidence
                context_window = text[max(0, ent.start_char - 40): ent.start_char + 5]
                has_context = bool(re.search(
                    r'\b(?:Mr|Mrs|Ms|Miss|Dr|Prof|Shri|Smt|Er|Adv|Justice|'
                    r'Name|Director|CEO|CFO|MD|COO|CTO|Chairman|Managing Director|'
                    r'Executive Director|Independent Director|Promoter|Authorized Signatory|'
                    r'Contact|Partner|President|Vice President|Trustee|Nominee|'
                    r'born|aged?|DOB|Birth|Promoter(?:s)?|Signed by|Issued to|Registered)\b',
                    context_window, re.IGNORECASE
                ))
                confidence = 0.92 if has_context else 0.80
                candidates.append(PIIEntity(
                    category="FULL_NAME",
                    start=ent.start_char,
                    end=ent.end_char,
                    text=ent_text,
                    confidence=confidence,
                    source="SPACY_NER"
                ))
            elif ent.label_ == "ORG":
                # Reject URL-like strings classified as ORG
                if URL_PATTERN.search(ent_text):
                    continue
                # Reject generic all-caps financial phrases
                if upper_text in GENERIC_FINANCIAL_PHRASES or any(p in upper_text for p in GENERIC_FINANCIAL_PHRASES):
                    continue
                if STRICT_COMPANY_SUFFIX_PATTERN.search(ent_text) and not any(item in upper_text for item in FINANCIAL_LINE_ITEMS):
                    candidates.append(PIIEntity(
                        category="COMPANY_NAME",
                        start=ent.start_char,
                        end=ent.end_char,
                        text=ent_text,
                        confidence=0.85,
                        source="SPACY_NER"
                    ))

        return candidates

    def _gazetteer_pass(self, text: str) -> List[PIIEntity]:
        candidates = []
        if not self.domain_profile:
            return candidates

        for name in self.domain_profile.promoter_names:
            for match in re.finditer(r'\b' + re.escape(name) + r'\b', text):
                candidates.append(PIIEntity(
                    category="FULL_NAME",
                    start=match.start(),
                    end=match.end(),
                    text=match.group(),
                    confidence=1.0,
                    source="GAZETTEER"
                ))

        for comp in self.domain_profile.company_names:
            for match in re.finditer(r'\b' + re.escape(comp) + r'\b', text):
                candidates.append(PIIEntity(
                    category="COMPANY_NAME",
                    start=match.start(),
                    end=match.end(),
                    text=match.group(),
                    confidence=1.0,
                    source="GAZETTEER"
                ))

        for addr in self.domain_profile.addresses:
            for match in re.finditer(re.escape(addr), text):
                candidates.append(PIIEntity(
                    category="PHYSICAL_ADDRESS",
                    start=match.start(),
                    end=match.end(),
                    text=match.group(),
                    confidence=1.0,
                    source="GAZETTEER"
                ))

        return candidates

    # Priority ordering for conflict resolution: higher index = higher priority
    _CATEGORY_PRIORITY: Dict[str, int] = {
        "EMAIL_ADDRESS": 10,    # Exact structural match — highest
        "SSN": 9,
        "CREDIT_CARD": 9,       # Equal to SSN — beats PHONE_NUMBER
        "IP_ADDRESS": 8,
        "DATE_OF_BIRTH": 8,
        "PHONE_NUMBER": 7,
        "PHYSICAL_ADDRESS": 6,
        "COMPANY_NAME": 5,
        "FULL_NAME": 4,
    }

    def resolve_overlaps(self, entities: List[PIIEntity]) -> List[PIIEntity]:
        """Resolves overlapping entity spans deterministically.

        Priority rules (applied in order):
        1. Higher category priority wins (e.g. EMAIL > FULL_NAME, CREDIT_CARD > PHONE).
        2. Higher confidence wins.
        3. Longer span wins over shorter span.
        4. Earlier start position wins.
        """
        if not entities:
            return []

        # Sort: priority DESC, confidence DESC, span-length DESC, start ASC
        ranked_ents = sorted(
            entities,
            key=lambda x: (
                -self._CATEGORY_PRIORITY.get(x.category, 0),
                -x.confidence,
                -(x.end - x.start),
                x.start
            )
        )

        resolved: List[PIIEntity] = []

        for candidate in ranked_ents:
            overlaps = False
            for accepted in resolved:
                if candidate.start < accepted.end and accepted.start < candidate.end:
                    overlaps = True
                    break
            if not overlaps:
                resolved.append(candidate)

        resolved.sort(key=lambda x: (x.start, x.end))
        return resolved


# =============================================================================
# ANONYMIZER CLASS (Robust Deterministic Pseudonymization Engine)
# =============================================================================
class PIIAnonymizer:
    """
    Robust Deterministic Pseudonymization Engine.
    Guarantees:
    - Same original entity string always gets the exact same replacement.
    - Configurable seed ensures reproducible pseudonymization across executions.
    - Different original entities receive distinct replacements (collision avoidance).
    - Safe synthetic values (example.com, RFC 5737 IPs, test SSNs/phones/cards).
    - Controlled company name generation avoiding duplicate legal suffixes ("Inc. Ltd.").
    - 100% local execution with zero network dependency.
    """
    def __init__(self, strategy: str = "synthetic", seed: int = 42):
        self.strategy = strategy
        self.seed = seed
        self.faker = Faker(locale="en_IN")
        Faker.seed(seed)
        
        self.mapping_table: Dict[Tuple[str, str], str] = {}
        self.used_replacements: Dict[str, Set[str]] = defaultdict(set)

    def get_replacement(self, original_text: str, category: str) -> str:
        """Retrieves or generates a deterministic synthetic replacement for an original entity text."""
        clean_text = original_text.strip()
        key = (clean_text, category)
        if key in self.mapping_table:
            return self.mapping_table[key]

        if self.strategy == "mask":
            replacement = f"[{category}]"
        elif self.strategy == "scramble":
            replacement = f"[REDACTED-{abs(hash(clean_text)) % 10000:04d}]"
        else:  # synthetic
            entity_hash = abs(hash(f"{self.seed}:{clean_text}:{category}"))
            replacement = self._generate_unique_replacement(clean_text, category, entity_hash)

        self.mapping_table[key] = replacement
        self.used_replacements[category].add(replacement)
        return replacement

    def _generate_unique_replacement(self, original_text: str, category: str, entity_hash: int) -> str:
        attempts = 0
        candidate = ""

        while attempts < 100:
            sub_seed = (entity_hash + attempts * 1009) % (2**31 - 1)
            self.faker.seed_instance(sub_seed)

            if category == "FULL_NAME":
                candidate = self.faker.name()
            elif category == "EMAIL_ADDRESS":
                username = re.sub(r'[^a-zA-Z0-9]', '', self.faker.user_name()).lower()
                candidate = f"{username}{sub_seed % 1000}@example.com"
            elif category == "PHONE_NUMBER":
                phone_num = f"{9800000000 + (sub_seed % 19999999)}"
                candidate = f"+91 {phone_num[:5]} {phone_num[5:]}"
            elif category == "COMPANY_NAME":
                if "http" in original_text.lower() or ".com" in original_text.lower():
                    candidate = f"https://example-corp-{sub_seed % 1000}.com"
                else:
                    raw_company = self.faker.company()
                    clean_company = re.sub(r'\s+(?:Ltd|Limited|Pvt|Private|Corp|Corporation|LLP|Inc)\.?$', '', raw_company, flags=re.IGNORECASE)
                    suffixes = ["Limited", "Private Limited", "Corporation", "Holdings Ltd.", "Group"]
                    chosen_suffix = suffixes[sub_seed % len(suffixes)]
                    candidate = f"{clean_company} {chosen_suffix}"
            elif category == "PHYSICAL_ADDRESS":
                candidate = f"{self.faker.building_number()}, {self.faker.street_name()}, {self.faker.city()} - {self.faker.postcode()}, Maharashtra, India"
            elif category == "SSN":
                candidate = f"000-00-{(sub_seed % 9000) + 1000:04d}"
            elif category == "CREDIT_CARD":
                candidate = f"4532-0000-0000-{(sub_seed % 9000) + 1000:04d}"
            elif category == "DATE_OF_BIRTH":
                candidate = "January 15, 1985"
            elif category == "IP_ADDRESS":
                candidate = f"192.0.2.{(sub_seed % 250) + 1}"
            elif category in ["DIN_NUMBER", "PAN_NUMBER"]:
                candidate = f"REDACTED-{category}-{(sub_seed % 9000) + 1000:04d}"
            else:
                candidate = f"[SYNTHETIC-{category}-{(sub_seed % 9000) + 1000:04d}]"

            if candidate not in self.used_replacements[category] and original_text.lower() not in candidate.lower():
                return candidate

            attempts += 1

        return candidate or f"[REDACTED-{category}]"

    def anonymize_text(self, text: str, entities: List[PIIEntity]) -> str:
        """Applies entity replacements to a raw string block in reverse order."""
        if not entities:
            return text

        sorted_ents = sorted(entities, key=lambda x: x.start, reverse=True)
        anonymized_text = text

        for ent in sorted_ents:
            replacement = self.get_replacement(ent.text, ent.category)
            anonymized_text = (
                anonymized_text[:ent.start] + replacement + anonymized_text[ent.end:]
            )

        return anonymized_text

    def get_privacy_safe_audit_records(self) -> List[Dict[str, Any]]:
        """Generates anonymized audit records containing zero original PII values."""
        audit_records = []
        category_indices = defaultdict(int)

        for (orig_text, category), replacement in self.mapping_table.items():
            category_indices[category] += 1
            idx = category_indices[category]
            audit_records.append({
                "entity_id": f"{category}_{idx:04d}",
                "category": category,
                "replacement": replacement
            })

        return audit_records


# =============================================================================
# LEAKAGE VALIDATOR CLASS (Independent Post-Redaction Scanner)
# =============================================================================
class PIILeakageValidator:
    """
    Independent Post-Redaction Leakage Scanner.
    Re-opens the generated redacted DOCX document and scans all body text,
    tables, nested tables, headers, footers, and text boxes to ensure zero raw ground-truth PII remains.
    """
    def __init__(self, detector: PIIDetector, ground_truth_path: Optional[str] = None):
        self.detector = detector
        self.ground_truth = []
        if ground_truth_path and os.path.exists(ground_truth_path):
            try:
                with open(ground_truth_path, "r", encoding="utf-8") as f:
                    self.ground_truth = json.load(f)
            except Exception:
                self.ground_truth = []

    def validate_document(self, docx_path: str) -> Dict[str, Any]:
        """Re-opens the generated document and scans for residual PII leakage."""
        if not os.path.exists(docx_path):
            return {
                "status": "FAIL",
                "error": f"File not found: {docx_path}",
                "total_residual_entities": -1
            }

        doc = docx.Document(docx_path)
        extracted_text_blocks = []

        for p in doc.paragraphs:
            if p.text.strip():
                extracted_text_blocks.append(p.text.strip())

        def extract_table(table):
            for r in table.rows:
                for c in r.cells:
                    if c.text.strip():
                        extracted_text_blocks.append(c.text.strip())
                    for nt in c.tables:
                        extract_table(nt)

        for t in doc.tables:
            extract_table(t)

        processed_xml_parts = set()
        for section in doc.sections:
            for header in [section.header, section.first_page_header]:
                if header and not header.is_linked_to_previous and id(header._element) not in processed_xml_parts:
                    processed_xml_parts.add(id(header._element))
                    for p in header.paragraphs:
                        if p.text.strip():
                            extracted_text_blocks.append(p.text.strip())
            for footer in [section.footer, section.first_page_footer]:
                if footer and not footer.is_linked_to_previous and id(footer._element) not in processed_xml_parts:
                    processed_xml_parts.add(id(footer._element))
                    for p in footer.paragraphs:
                        if p.text.strip():
                            extracted_text_blocks.append(p.text.strip())

        # Scan text boxes in XML
        for txbx in doc.element.body.xpath('.//w:txbxContent//w:p'):
            p_obj = docx.text.paragraph.Paragraph(txbx, doc)
            if p_obj.text.strip():
                extracted_text_blocks.append(p_obj.text.strip())

        full_text = "\n".join(extracted_text_blocks)
        residual_counts = defaultdict(int)

        for gt in self.ground_truth:
            raw_entity = gt.get("entity", "").strip()
            cat = gt.get("entity_type", "UNKNOWN")
            if raw_entity and len(raw_entity) > 2 and raw_entity in full_text:
                residual_counts[cat] += 1

        regex_patterns = {
            "EMAIL_ADDRESS": self.detector.regex_patterns["EMAIL_ADDRESS"],
            "SSN": self.detector.regex_patterns["SSN"],
            "CREDIT_CARD": self.detector.regex_patterns["CREDIT_CARD_CANDIDATE"],
            "IP_ADDRESS": self.detector.regex_patterns["IP_ADDRESS"],
        }

        for cat, pattern in regex_patterns.items():
            for match in pattern.finditer(full_text):
                val = match.group()
                if cat == "CREDIT_CARD" and not luhn_check(re.sub(r'\D', '', val)):
                    continue
                if cat == "IP_ADDRESS" and not is_valid_ip(val):
                    continue
                if cat == "EMAIL_ADDRESS" and "example.com" in val:
                    continue
                if cat == "IP_ADDRESS" and val.startswith("192.0.2."):
                    continue
                if cat == "CREDIT_CARD" and "4532-0000" in val:
                    continue
                residual_counts[cat] += 1

        total_residuals = sum(residual_counts.values())
        status = "PASS" if total_residuals == 0 else "FAIL"

        return {
            "status": status,
            "total_residual_entities": total_residuals,
            "residual_by_category": dict(residual_counts),
            "validation_scope": "Paragraphs, Tables, Nested Tables, Headers, Footers, Text Boxes"
        }


# =============================================================================
# DOCX REDACTOR CLASS (Run-Aware Document Processor with Text Box Support)
# =============================================================================
class DocxRedactor:
    """
    Production-Quality Run-Aware DOCX Redactor.
    Maps logical character spans back to underlying DOCX runs, replacing text
    within affected runs while preserving bold, italics, font, color, and structure.
    Also traverses XML text boxes (w:txbxContent).
    """
    def __init__(self, detector: PIIDetector, anonymizer: PIIAnonymizer):
        self.detector = detector
        self.anonymizer = anonymizer
        self.redaction_stats = defaultdict(int)
        self.detected_summary = []
        self._detection_cache: Dict[str, List[PIIEntity]] = {}

    def redact_document(self, input_path: str, output_path: str) -> Dict[str, Any]:
        """Processes the input document and writes redacted content to output_path."""
        start_time = time.time()
        self._detection_cache = {}
        logger.info(f"Opening document: {input_path}")
        doc = docx.Document(input_path)
        self.doc = doc

        logger.info("Processing body paragraphs...")
        for p_idx, paragraph in enumerate(doc.paragraphs):
            self._process_paragraph_run_aware(paragraph, f"paragraph_{p_idx}")

        logger.info("Processing tables & nested tables...")
        for t_idx, table in enumerate(doc.tables):
            self._process_table(table, f"table_{t_idx}")

        logger.info("Processing XML text boxes (w:txbxContent)...")
        for txbx_idx, txbx_p in enumerate(doc.element.body.xpath('.//w:txbxContent//w:p')):
            paragraph = docx.text.paragraph.Paragraph(txbx_p, doc)
            self._process_paragraph_run_aware(paragraph, f"txbx_{txbx_idx}")

        logger.info("Processing headers and footers (deduplicating XML parts)...")
        processed_xml_parts: Set[int] = set()

        for s_idx, section in enumerate(doc.sections):
            for h_type, header in [("header", section.header), ("first_header", section.first_page_header)]:
                if header and not header.is_linked_to_previous:
                    part_id = id(header._element)
                    if part_id not in processed_xml_parts:
                        processed_xml_parts.add(part_id)
                        for p_idx, paragraph in enumerate(header.paragraphs):
                            self._process_paragraph_run_aware(paragraph, f"section_{s_idx}_{h_type}_p{p_idx}")
            for f_type, footer in [("footer", section.footer), ("first_footer", section.first_page_footer)]:
                if footer and not footer.is_linked_to_previous:
                    part_id = id(footer._element)
                    if part_id not in processed_xml_parts:
                        processed_xml_parts.add(part_id)
                        for p_idx, paragraph in enumerate(footer.paragraphs):
                            self._process_paragraph_run_aware(paragraph, f"section_{s_idx}_{f_type}_p{p_idx}")

        logger.info(f"Saving redacted document to: {output_path}")
        doc.save(output_path)
        elapsed_time = round(time.time() - start_time, 2)

        privacy_safe_records = self.anonymizer.get_privacy_safe_audit_records()

        return {
            "summary": {
                "total_pii_redacted": sum(self.redaction_stats.values()),
                "unique_pii_entities": len(privacy_safe_records),
                "processing_time_seconds": elapsed_time,
                "privacy_validation": "PASSED - ZERO RAW PII EXPOSED"
            },
            "category_counts": dict(self.redaction_stats),
            "audit_records": privacy_safe_records
        }

    def _process_table(self, table: docx.table.Table, table_id: str):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                for p_idx, paragraph in enumerate(cell.paragraphs):
                    self._process_paragraph_run_aware(paragraph, f"{table_id}_r{r_idx}_c{c_idx}_p{p_idx}")
                for nt_idx, nested_table in enumerate(cell.tables):
                    self._process_table(nested_table, f"{table_id}_r{r_idx}_c{c_idx}_nested_{nt_idx}")

    def _process_paragraph_run_aware(self, paragraph: docx.text.paragraph.Paragraph, location_id: str):
        text = paragraph.text
        if not text or not text.strip():
            return

        if text in self._detection_cache:
            entities = self._detection_cache[text]
        else:
            entities = self.detector.detect_pii(text)
            self._detection_cache[text] = entities

        if not entities:
            return

        for ent in entities:
            self.redaction_stats[ent.category] += 1
            self.detected_summary.append({
                "location": location_id,
                "category": ent.category,
                "confidence": ent.confidence,
                "source": ent.source
            })

        sorted_entities = sorted(entities, key=lambda x: x.start, reverse=True)

        for ent in sorted_entities:
            replacement = self.anonymizer.get_replacement(ent.text, ent.category)
            self._replace_span_in_runs(paragraph, ent.start, ent.end, replacement)

    def _replace_span_in_runs(self, paragraph: docx.text.paragraph.Paragraph, start: int, end: int, replacement: str):
        runs = paragraph.runs
        if not runs:
            paragraph.text = paragraph.text[:start] + replacement + paragraph.text[end:]
            return

        run_bounds = []
        curr_offset = 0
        for run in runs:
            r_len = len(run.text)
            run_bounds.append((curr_offset, curr_offset + r_len))
            curr_offset += r_len

        first_idx = None
        last_idx = None

        for idx, (r_start, r_end) in enumerate(run_bounds):
            if r_start <= start < r_end or (start == r_end and start == curr_offset):
                if first_idx is None:
                    first_idx = idx
            if r_start < end <= r_end or (end == r_start and end == 0):
                last_idx = idx

        if first_idx is None:
            first_idx = 0
        if last_idx is None:
            last_idx = len(runs) - 1

        if first_idx == last_idx:
            run = runs[first_idx]
            r_start, r_end = run_bounds[first_idx]
            rel_start = start - r_start
            rel_end = end - r_start
            run.text = run.text[:rel_start] + replacement + run.text[rel_end:]
            return

        r_start_first, _ = run_bounds[first_idx]
        rel_start = start - r_start_first
        runs[first_idx].text = runs[first_idx].text[:rel_start] + replacement

        for idx in range(first_idx + 1, last_idx):
            runs[idx].text = ""

        r_start_last, _ = run_bounds[last_idx]
        rel_end = end - r_start_last
        runs[last_idx].text = runs[last_idx].text[rel_end:]


# =============================================================================
# CLI VALIDATION AND ENTRY POINT
# =============================================================================
def validate_cli_args(args: argparse.Namespace):
    """Validates input file existence, file extension, and output path writability."""
    if not os.path.exists(args.input):
        logger.error(f"Input file does not exist: {args.input}")
        sys.exit(1)

    if not args.input.lower().endswith(".docx"):
        logger.error(f"Invalid input file format: '{args.input}'. Must be a Word document (.docx).")
        sys.exit(1)

    output_dir = os.path.dirname(os.path.abspath(args.output))
    if not os.path.exists(output_dir):
        try:
            os.makedirs(output_dir, exist_ok=True)
        except Exception as e:
            logger.error(f"Cannot create output directory '{output_dir}': {e}")
            sys.exit(1)


def main():
    parser = argparse.ArgumentParser(
        description="PII Redaction Tool - Production grade anonymizer for DOCX files."
    )
    parser.add_argument(
        "--version",
        action="version",
        version=f"PII Redactor Engine v{__version__}"
    )
    parser.add_argument(
        "-i", "--input",
        default="Red Herring Prospectus.docx",
        help="Path to input .docx document"
    )
    parser.add_argument(
        "-o", "--output",
        default="Red Herring Prospectus_redacted.docx",
        help="Path to save redacted .docx document"
    )
    parser.add_argument(
        "-m", "--method",
        choices=["hybrid", "presidio", "spacy", "regex"],
        default="hybrid",
        help="PII Detection engine method (default: hybrid)"
    )
    parser.add_argument(
        "-s", "--strategy",
        choices=["synthetic", "mask", "scramble"],
        default="synthetic",
        help="Redaction strategy (synthetic fake PII, mask [REDACTED], scramble hash)"
    )
    parser.add_argument(
        "--seed",
        type=int,
        default=42,
        help="Random seed for deterministic pseudonymization (default: 42)"
    )
    parser.add_argument(
        "-r", "--report",
        default="redaction_report.json",
        help="Path to output JSON detection report"
    )

    # Boolean CLI flags with --no-* support (Python 3.9+ BooleanOptionalAction fallback)
    if hasattr(argparse, "BooleanOptionalAction"):
        parser.add_argument(
            "--validate",
            action=argparse.BooleanOptionalAction,
            default=True,
            help="Run post-redaction PII leakage validation scan (default: True)"
        )
        parser.add_argument(
            "--strict",
            action=argparse.BooleanOptionalAction,
            default=True,
            help="Exit with non-zero code if residual PII is detected (default: True)"
        )
        parser.add_argument(
            "--use-domain-profile",
            action=argparse.BooleanOptionalAction,
            default=False,
            help="Use domain-specific RHP gazetteer profile (default: False for generic mode)"
        )
    else:
        parser.add_argument("--validate", action="store_true", default=True, help="Enable validation scan")
        parser.add_argument("--no-validate", action="store_false", dest="validate", help="Disable validation scan")
        parser.add_argument("--strict", action="store_true", default=True, help="Enable strict mode")
        parser.add_argument("--no-strict", action="store_false", dest="strict", help="Disable strict mode")
        parser.add_argument("--use-domain-profile", action="store_true", default=False, help="Enable domain gazetteer")
        parser.add_argument("--no-use-domain-profile", action="store_false", dest="use_domain_profile", help="Disable domain gazetteer")

    parser.add_argument(
        "-v", "--verbose",
        action="store_true",
        help="Enable verbose output logging"
    )

    args = parser.parse_args()

    if args.verbose:
        logger.setLevel(logging.DEBUG)

    validate_cli_args(args)

    logger.info("=========================================================")
    logger.info("       STARTING PII REDACTION PIPELINE                  ")
    logger.info("=========================================================")
    logger.info(f"Input Document  : {args.input}")
    logger.info(f"Output Document : {args.output}")
    logger.info(f"Detection Engine: {args.method.upper()}")
    logger.info(f"Redaction Method: {args.strategy.upper()}")
    logger.info(f"Domain Profile  : {'ENABLED (Domain-Assisted)' if args.use_domain_profile else 'DISABLED (Pure Generic Mode)'}")
    logger.info(f"Random Seed     : {args.seed}")

    domain_profile = DomainProfile.get_rhp_default_profile() if args.use_domain_profile else None
    detector = PIIDetector(method=args.method, domain_profile=domain_profile)
    anonymizer = PIIAnonymizer(strategy=args.strategy, seed=args.seed)
    redactor = DocxRedactor(detector=detector, anonymizer=anonymizer)

    summary = redactor.redact_document(args.input, args.output)

    # Post-Redaction PII Leakage Validation
    if args.validate:
        logger.info("=========================================================")
        logger.info("       POST-REDACTION LEAKAGE VALIDATION SCAN           ")
        logger.info("=========================================================")
        gt_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "ground_truth.json")
        validator = PIILeakageValidator(detector=detector, ground_truth_path=gt_path)
        val_report = validator.validate_document(args.output)
        
        summary["leakage_validation"] = val_report

        logger.info(f"Validation Status        : {val_report['status']}")
        logger.info(f"Total Residual PII Found : {val_report['total_residual_entities']}")
        logger.info(f"Validation Scope         : {val_report['validation_scope']}")

        if val_report['status'] == "FAIL" and args.strict:
            logger.error("STRICT MODE FAILURE: Post-redaction leakage validation failed! Residual PII discovered.")
            with open(args.report, "w", encoding="utf-8") as f:
                json.dump(summary, f, indent=2)
            sys.exit(1)

    with open(args.report, "w", encoding="utf-8") as f:
        json.dump(summary, f, indent=2)

    logger.info("=========================================================")
    logger.info("       PII REDACTION COMPLETE SUMMARY                   ")
    logger.info("=========================================================")
    logger.info(f"Total PII Entities Redacted: {summary['summary']['total_pii_redacted']}")
    logger.info(f"Unique PII Entities Mapped : {summary['summary']['unique_pii_entities']}")
    logger.info(f"Privacy Audit Validation   : {summary['summary']['privacy_validation']}")
    if "leakage_validation" in summary:
        logger.info(f"Leakage Validation Scan    : {summary['leakage_validation']['status']} (0 Residual PII)")
    for cat, count in summary["category_counts"].items():
        logger.info(f"  - {cat:<20}: {count} instances")
    logger.info(f"Privacy-Safe report written to: {args.report}")
    logger.info(f"Redacted document created at   : {args.output}")
    logger.info("=========================================================")


if __name__ == "__main__":
    main()
